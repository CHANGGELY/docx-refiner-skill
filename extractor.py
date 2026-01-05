import docx
import os

def extract_all_text(file_path):
    """
    从 .docx 文件中提取所有文本段落。
    """
    doc = docx.Document(file_path)
    # 提取所有段落的文本，过滤掉纯空白行
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    return paragraphs

def extract_full_document(file_path, output_dir):
    """
    遍历文档，返回一个按顺序排列的列表，包含文本和图片。
    格式: [{"type": "text", "content": "..."}, {"type": "image", "content": "path/to/img"}]
    """
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
        
    doc = docx.Document(file_path)
    content_list = []
    image_counter = 0
    
    # 辅助函数：保存图片并返回路径
    def save_image(part, counter):
        ext = part.content_type.split("/")[-1]
        image_name = f"image_{counter}.{ext}"
        image_path = os.path.join(output_dir, image_name)
        with open(image_path, "wb") as f:
            f.write(part.blob)
        return image_path

    # 遍历段落
    for paragraph in doc.paragraphs:
        # 先处理文本（如果段落有文字）
        if paragraph.text.strip():
            content_list.append({"type": "text", "content": paragraph.text.strip()})
        
        # 处理段落内的图片 (In-line shapes)
        # 在 docx-python 中，图片通常通过 paragraph.runs 里的 element 找
        for run in paragraph.runs:
            # 查找 run 里的 drawing 元素
            # 这是一个简化的探测方式，docx-python 的 API 在图片位置上比较底层
            if "drawing" in run._element.xml:
                # 寻找关联的 blipId
                import re
                blip_ids = re.findall(r'r:embed="([^"]+)"', run._element.xml)
                for blip_id in blip_ids:
                    if blip_id in paragraph.part.rels:
                        target_part = paragraph.part.rels[blip_id].target_part
                        if "image" in target_part.content_type:
                            img_path = save_image(target_part, image_counter)
                            content_list.append({"type": "image", "content": img_path})
                            image_counter += 1
                            
    return content_list

if __name__ == "__main__":
    # 简单测试
    import sys
    if len(sys.argv) > 1:
        res = extract_all_text(sys.argv[1])
        for i, p in enumerate(res[:5]):
            print(f"[{i}] {p}")
