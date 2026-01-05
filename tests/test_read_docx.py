import docx
import os
import sys

def test_read_docx():
    # 从环境变量获取配置，默认值为硬编码值
    file_path = os.getenv("TEST_DOCX_FILE", "code之路.docx")
    if not os.path.exists(file_path):
        print(f"Error: File not found at {file_path}")
        sys.exit(1)
    
    try:
        doc = docx.Document(file_path)
        print(f"Successfully opened {file_path}")
        print(f"Paragraph count: {len(doc.paragraphs)}")
        if len(doc.paragraphs) == 0:
            print("Warning: Document has 0 paragraphs.")
    except Exception as e:
        print(f"Failed to read document: {e}")
        sys.exit(1)

if __name__ == "__main__":
    test_read_docx()
